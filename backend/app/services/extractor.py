from pathlib import Path
import pdfplumber
from docx import Document

def extract_text(file_path):
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf(path)

    if extension == ".docx":
        return extract_docx(path)

    if extension in {".txt", ".tex"}:
        return extract_plain_text(path)

    raise ValueError("Unsupported file type")

def extract_pdf(path):
    chunks = []

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text.strip())

    extracted = "\n".join(chunks).strip()

    if not extracted:
        raise ValueError("No text could be extracted from this PDF. It may be scanned or image-based.")

    return normalize_text(extracted)

def extract_docx(path):
    document = Document(str(path))
    chunks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)

    extracted = "\n".join(chunks).strip()

    if not extracted:
        raise ValueError("No text could be extracted from this DOCX file.")

    return normalize_text(extracted)

def extract_plain_text(path):
    try:
        extracted = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        extracted = path.read_text(encoding="latin-1")

    extracted = extracted.strip()

    if not extracted:
        raise ValueError("The uploaded text file is empty.")

    return normalize_text(extracted)

def normalize_text(text):
    replacements = {
        "“": "\"",
        "”": "\"",
        "‘": "'",
        "’": "'",
        "—": "-",
        "–": "-",
        "\xa0": " "
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    return "\n".join(lines)
